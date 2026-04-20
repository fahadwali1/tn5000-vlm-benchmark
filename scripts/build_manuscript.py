#!/usr/bin/env python3
"""
Build FINAL npj Digital Medicine manuscript + Supplementary Information.
- Full zero-shot table (all 36 configs: 12 models x 3 prompts)
- 4 final polish edits applied
- Supplementary sections included
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)
font.color.rgb = RGBColor(0, 0, 0)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.5

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

FIGS = '/Users/fahadwali/Downloads/TN5000/figures'

def add_heading_n(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Arial'
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_p(text, bold=False, italic=False, size=None, align=None, sa=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if align: p.alignment = align
    if sa is not None: p.paragraph_format.space_after = Pt(sa)
    return p

def set_cell_shade(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def fmt_table(table, hdr_color='2F5496'):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="6" w:space="0" w:color="2F5496"/>'
        '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="2F5496"/>'
        '  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="D9E2F3"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    for cell in table.rows[0].cells:
        set_cell_shade(cell, hdr_color)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(7.5)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(7.5)
            if i % 2 == 0:
                set_cell_shade(cell, 'F2F7FC')

def add_fig(path, width, caption):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_p(caption, italic=True, size=9, sa=14)


# ============================================================
# TITLE PAGE
# ============================================================
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run(
    'Parameter-Efficient Fine-Tuning Corrects Specificity Collapse in '
    'Zero-Shot Vision-Language Models for Thyroid Ultrasound Classification'
)
tr.font.name = 'Arial'; tr.font.size = Pt(16); tr.bold = True
tp.paragraph_format.space_after = Pt(12)

add_p('Fahad Wali Ahmed\u00b9*, Fauwad Wali Ahmed\u00b2, Belal Moftah\u00b9\u02cc\u00b3',
      size=11, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)

# Affiliations
add_p('\u00b9King Faisal Specialist Hospital and Research Centre\u2013Madinah, Saudi Arabia',
      size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
add_p('\u00b2School of Business, University of Houston\u2013Clear Lake, Houston, TX, USA',
      size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
add_p('\u00b3Medical Physics Unit, Gerald Bronfman Department of Oncology, '
      'McGill University, Montr\u00e9al, QC, Canada',
      size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)

# ORCID line
add_p('ORCID: F.W.A. 0000-0003-4100-1571; F.A. 0009-0009-6805-2986; '
      'B.M. 0000-0003-0509-9831',
      size=8, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)

# Corresponding author
add_p('*Corresponding author: Fahad Wali Ahmed, fahadwali@yahoo.com',
      size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=18)

# ============================================================
# ABSTRACT (unstructured, <200 words)
# ============================================================
add_heading_n('Abstract', 1)

add_p(
    'Vision-language models (VLMs) offer zero-shot medical image interpretation, '
    'but their diagnostic reliability remains unknown. We benchmarked 18 models on the '
    'TN5000 thyroid ultrasound dataset (n\u2009=\u20091,000 biopsy-confirmed test images) across three '
    'paradigms: 12 VLMs evaluated zero-shot under three prompt strategies (baseline, chain-of-thought, '
    'and TI-RADS guided), three open-weight VLMs fine-tuned with quantised low-rank adaptation (QLoRA), '
    'and three supervised classifiers. Zero-shot VLMs achieved 62.3-74.5% accuracy but exhibited '
    'uniformly poor specificity (median 13.8%), reflecting systematic malignancy-prediction bias. '
    'QLoRA fine-tuning of LLaVA-NeXT 7B corrected this deficit, achieving 87.1% accuracy, 84.4% '
    'specificity, and AUC\u2009=\u20090.938; performance numerically similar to the best supervised model '
    '(Swin-T: 85.6%, AUC\u2009=\u20090.937; p\u2009=\u20090.33) while modifying only 0.2% of parameters. '
    'Notably, the supervised classifier achieved this comparable performance with substantially greater '
    'training efficiency (5.3 minutes vs. 179.7 minutes for LLaVA-NeXT), a 30-fold difference with '
    'implications for deployment feasibility. At 5% clinical prevalence, zero-shot VLMs would generate '
    '17.4 false positives per true positive, whereas fine-tuned LLaVA-NeXT reduced this five-fold to 3.4. '
    'Fine-tuning of Gemma 4 27B failed '
    'to converge, indicating model size does not predict adaptation success. On this benchmark, '
    'zero-shot VLMs were clinically unreliable for thyroid nodule classification; parameter-efficient '
    'fine-tuning can correct this in selected architectures. Multi-centre external validation is needed.',
    sa=6
)

kwp = doc.add_paragraph()
kwr1 = kwp.add_run('Keywords: '); kwr1.bold = True; kwr1.font.name = 'Arial'; kwr1.font.size = Pt(9)
kwr2 = kwp.add_run(
    'vision-language models; thyroid nodule classification; ultrasound; '
    'QLoRA fine-tuning; deep learning; parameter-efficient fine-tuning'
); kwr2.font.name = 'Arial'; kwr2.font.size = Pt(9)
kwp.paragraph_format.space_after = Pt(18)

# ============================================================
# INTRODUCTION
# ============================================================
add_heading_n('Introduction', 1)

add_p(
    'Thyroid nodules are among the most common incidental findings in clinical imaging, '
    'with prevalence rates of 20-76% in the general population detected by high-resolution '
    'ultrasonography\u00b9. Against a backdrop of rising global thyroid cancer incidence\u00b2, '
    'approximately 5-15% of nodules harbour malignancy, necessitating systematic risk '
    'stratification to identify those warranting fine-needle aspiration biopsy\u00b3. The American College of Radiology Thyroid Imaging, '
    'Reporting and Data System (ACR TI-RADS) provides a standardised framework for this purpose\u2074, '
    'but remains subject to substantial interobserver variability.'
)

add_p(
    'Deep learning approaches have demonstrated considerable promise for automated thyroid nodule '
    'classification on ultrasound, with convolutional neural networks (CNNs) and vision transformers '
    'achieving diagnostic performance approaching that of experienced radiologists\u2075\u02cc\u2076\u02cc\u2077\u02cc\u2078\u02cc\u2079. '
    'These supervised models, however, require large annotated training datasets and are inherently '
    'limited to the specific classification task for which they were trained.'
)

add_p(
    'Concurrently, vision-language models (VLMs), multimodal architectures that jointly process '
    'images and natural language, have emerged as a potentially transformative technology for '
    'medical imaging\u00b9\u2070\u02cc\u00b9\u00b9. Unlike task-specific classifiers, VLMs can interpret images '
    'through natural language instructions, offering the theoretical advantages of zero-shot '
    'generalisation, explainable outputs, and flexible task adaptation. Recent studies have evaluated '
    'VLMs across radiology\u00b9\u00b2\u02cc\u00b9\u00b3, dermatology\u00b9\u2074, paediatric imaging\u00b9\u2075, '
    'orthopedics\u00b9\u2076, and dentistry\u00b9\u2077, with mixed results.'
)

add_p(
    'However, the clinical utility of VLMs for standardised medical imaging benchmarks remains '
    'poorly characterised. First, the zero-shot performance ceiling has not been systematically '
    'established. Second, the extent to which parameter-efficient fine-tuning (PEFT), particularly '
    'quantised low-rank adaptation (QLoRA), can bridge the gap between zero-shot and supervised '
    'performance is unknown. Recent work has applied LLM fine-tuning to clinical tasks including '
    'clinical letter generation in radiation oncology\u00b9\u2078, and broader foundations for biomedical LLMs have been '
    'reviewed\u00b9\u2079, including ophthalmology decision support\u00b2\u2070, but systematic evaluations on '
    'standardised imaging benchmarks are lacking. '
    'Third, fair comparisons between fine-tuned VLMs and purpose-built supervised classifiers on '
    'identical datasets are largely absent from the literature. Because thyroid imaging datasets '
    'are often class-imbalanced, evaluation requires metrics that do not overstate performance when '
    'models favour the majority class.'
)

add_p(
    'To address these gaps, we conducted a comprehensive three-tier benchmark study on the TN5000 '
    'dataset\u00b2\u00b9, a publicly available collection of 5,000 biopsy-confirmed thyroid ultrasound images. '
    'Our evaluation encompasses: (1) 12 state-of-the-art VLMs evaluated zero-shot across three prompt '
    'strategies; (2) QLoRA fine-tuning of three VLMs (LLaVA-NeXT 7B, Qwen2.5-VL-7B, '
    'and Gemma 4 27B); and (3) supervised training of three deep learning architectures '
    '(ResNet-50, EfficientNet-B0, Swin-T). This is, to our knowledge, among the first '
    'studies to systematically compare zero-shot, fine-tuned, and supervised paradigms for thyroid '
    'nodule classification on a single standardised benchmark.'
)

# ============================================================
# RESULTS
# ============================================================
add_heading_n('Results', 1)

# --- Zero-Shot ---
add_heading_n('Zero-shot VLM performance', 2)

add_p(
    'Across 36,000 image-level classifications (12 models \u00d7 3 prompts \u00d7 1,000 images), '
    'zero-shot VLMs achieved best-prompt accuracies ranging from 62.3% (Llama 4 Scout) to 74.5% '
    '(Qwen3-VL-235B). Table 1 presents the complete results for all 36 model-prompt configurations. '
    'The most striking finding was uniformly poor specificity across all VLMs. Median specificity across the '
    '12 models (using each model\u2019s best prompt) was '
    '13.8% (range 3.0-29.7%), indicating that VLMs classified the vast majority of benign nodules '
    'as malignant. Conversely, sensitivity was uniformly high (median 92.3%, range 75.8-99.5%) (Fig. 1).'
)

add_p(
    'This sensitivity-specificity asymmetry reflects a systematic malignancy-prediction bias: '
    'VLMs defaulted to predicting malignancy when uncertain, achieving high sensitivity at the cost '
    'of near-total specificity collapse. The best zero-shot model, Qwen3-VL-235B, achieved 74.5% '
    'accuracy with 98.2% sensitivity but only 10.0% specificity, yielding MCC \u2248 0.13, barely '
    'above random for a balanced metric.'
)

add_p(
    'Prompt engineering had inconsistent effects across models (Table 1). No single strategy was '
    'universally superior. The TI-RADS guided prompt improved GPT-4.1 Mini accuracy from 72.0% '
    '(baseline) to 73.5% but degraded Claude Sonnet 4 from 73.5% (baseline) to 46.0%. '
    'Chain-of-thought prompting benefited Gemini 2.0 Flash (71.3% to 73.0%) but caused Grok 4 '
    'to collapse (70.7% to 56.1%). Several models showed a specificity-sensitivity trade-off '
    'across prompts: GPT-4.1 achieved 62.5% specificity under baseline prompting but only 9.3% '
    'under TI-RADS, while accuracy rose from 64.0% to 72.5%.'
)

# TABLE 1: Full zero-shot (all 36 configurations)
add_p('Table 1 | Zero-shot VLM performance across all prompt strategies.', bold=True, size=9, sa=4)

# Full data: model, prompt, acc, sens, spec
zs_full = [
    # Qwen3-VL-235B
    ['Qwen3-VL-235B', 'Baseline', '0.745', '0.982', '0.100'],
    ['', 'CoT', '0.732', '0.997', '0.011'],
    ['', 'TI-RADS', '0.731', '1.000', '0.000'],
    # Qwen2.5-VL-72B
    ['Qwen2.5-VL-72B', 'Baseline', '0.738', '0.971', '0.108'],
    ['', 'CoT', '0.731', '0.985', '0.041'],
    ['', 'TI-RADS', '0.733', '0.989', '0.034'],
    # GPT-4.1 Mini
    ['GPT-4.1 Mini', 'Baseline', '0.720', '0.850', '0.365'],
    ['', 'CoT', '0.714', '0.870', '0.290'],
    ['', 'TI-RADS', '0.735', '0.988', '0.048'],
    # Claude Sonnet 4
    ['Claude Sonnet 4', 'Baseline', '0.735', '0.995', '0.030'],
    ['', 'CoT', '0.681', '0.865', '0.182'],
    ['', 'TI-RADS', '0.460', '0.358', '0.736'],
    # Gemini 2.0 Flash
    ['Gemini 2.0 Flash', 'Baseline', '0.713', '0.866', '0.297'],
    ['', 'CoT', '0.730', '0.906', '0.253'],
    ['', 'TI-RADS', '0.683', '0.818', '0.316'],
    # GPT-4.1
    ['GPT-4.1', 'Baseline', '0.640', '0.646', '0.625'],
    ['', 'CoT', '0.705', '0.867', '0.264'],
    ['', 'TI-RADS', '0.725', '0.958', '0.093'],
    # Gemini 2.5 Pro
    ['Gemini 2.5 Pro', 'Baseline', '0.724', '0.923', '0.186'],
    ['', 'CoT', '0.706', '0.895', '0.187'],
    ['', 'TI-RADS', '0.719', '0.940', '0.121'],
    # Claude Sonnet 4.6
    ['Claude Sonnet 4.6', 'Baseline', '0.539', '0.443', '0.799'],
    ['', 'CoT', '0.571', '0.557', '0.610'],
    ['', 'TI-RADS', '0.713', '0.926', '0.134'],
    # Grok 4
    ['Grok 4', 'Baseline', '0.707', '0.917', '0.138'],
    ['', 'CoT', '0.561', '0.568', '0.544'],
    ['', 'TI-RADS', '0.701', '0.928', '0.086'],
    # Gemini 2.5 Flash
    ['Gemini 2.5 Flash', 'Baseline', '0.661', '0.694', '0.573'],
    ['', 'CoT', '0.689', '0.840', '0.279'],
    ['', 'TI-RADS', '0.574', '0.620', '0.450'],
    # OpenAI o4-mini
    ['OpenAI o4-mini', 'Baseline', '0.602', '0.695', '0.351'],
    ['', 'CoT', '0.640', '0.771', '0.284'],
    ['', 'TI-RADS', '0.680', '0.887', '0.116'],
    # Llama 4 Scout
    ['Llama 4 Scout', 'Baseline', '0.459', '0.360', '0.729'],
    ['', 'CoT', '0.548', '0.528', '0.604'],
    ['', 'TI-RADS', '0.623', '0.758', '0.257'],
]

t1 = doc.add_table(rows=len(zs_full)+1, cols=5)
# Header
for j, h in enumerate(['Model', 'Prompt', 'Accuracy', 'Sensitivity', 'Specificity']):
    t1.rows[0].cells[j].text = h
# Data
for i, row_data in enumerate(zs_full):
    for j, val in enumerate(row_data):
        t1.rows[i+1].cells[j].text = val
fmt_table(t1)

# Left-align model + prompt columns; highlight best-prompt rows
for i in range(1, len(zs_full)+1):
    t1.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t1.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Bold the model name cells
    model_text = t1.rows[i].cells[0].text
    if model_text:  # non-empty = first row of a model group
        for run in t1.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

add_p(
    'Bold model names indicate first appearance. Best-performing prompt per model is highlighted '
    'in the Results text. All models were accessed through the OpenRouter API gateway.',
    italic=True, size=8, sa=12
)

# FIGURE 1: Sensitivity vs Specificity
add_fig(f'{FIGS}/Fig1_sensitivity_specificity.png', 5.5,
    'Fig. 1 | Sensitivity versus specificity across the benchmark: 13 best-prompt zero-shot '
    'configurations from 12 VLMs (blue circles), three QLoRA fine-tuned VLMs (red triangles), '
    'and three supervised baselines (green squares). Zero-shot models cluster in the upper-left '
    'quadrant with high sensitivity but uniformly poor specificity (shaded zone, <30%). '
    'Fine-tuned and supervised models achieve balanced performance in the upper-right quadrant. '
    'The Gemma 4 27B fine-tuning failure (grey cross) remains in the low-specificity zone.'
)

# --- Fine-Tuned ---
add_heading_n('Fine-tuned VLM performance', 2)

add_p(
    'QLoRA fine-tuning produced markedly different outcomes across the three VLMs tested (Table 2). '
    'All fine-tuning runs used a single random seed (seed\u2009=\u200942); variance across seeds was '
    'not assessed, and results should be interpreted as point estimates from single training runs.'
)

add_p(
    'LLaVA-NeXT 7B achieved the numerically highest performance in the benchmark: accuracy 87.1% '
    '(95% CI 84.9-89.0%), sensitivity 88.1% (95% CI 85.6-90.3%), specificity 84.4% '
    '(95% CI 79.6-88.2%), F1\u2009=\u20090.909, MCC\u2009=\u20090.693, and AUC-ROC\u2009=\u20090.938. '
    'Compared to the best zero-shot VLM, this represents a 12.6 percentage-point accuracy improvement '
    'and a 74.4 percentage-point improvement in specificity (from 10.0% to 84.4%). Training required '
    '179.7 minutes on a single A100-80GB GPU.'
)

add_p(
    'Qwen2.5-VL-7B showed substantial but more modest improvement: accuracy 80.2%, sensitivity 81.9%, '
    'specificity 75.5%, F1\u2009=\u20090.858, MCC\u2009=\u20090.539, and AUC-ROC\u2009=\u20090.862. '
    'Gemma 4 27B failed to benefit from fine-tuning, achieving only 62.5% accuracy with 14.5% '
    'specificity \u2014 performance comparable to zero-shot models. Despite its larger parameter count '
    '(27B vs. 7B), the model exhibited training instability with predictions collapsing toward the '
    'majority class under multiple hyperparameter configurations.'
)

# TABLE 2: Fine-tuned
add_p('Table 2 | Fine-tuned VLM performance (QLoRA).', bold=True, size=9, sa=4)
t2d = [
    ['Model', 'Accuracy', 'Sensitivity', 'Specificity', 'F1', 'MCC', 'AUC-ROC', 'Train time'],
    ['LLaVA-NeXT 7B', '0.871', '0.881', '0.844', '0.909', '0.693', '0.938', '179.7 min'],
    ['Qwen2.5-VL-7B', '0.802', '0.819', '0.755', '0.858', '0.539', '0.862', '161.7 min'],
    ['Gemma 4 27B*', '0.625', '0.802', '0.145', '0.758', 'N/A', 'N/A', '~300 min'],
]
t2 = doc.add_table(rows=4, cols=8)
for i, rd in enumerate(t2d):
    for j, v in enumerate(rd):
        t2.rows[i].cells[j].text = v
fmt_table(t2)
for i in range(1, 4):
    t2.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
add_p('*Gemma 4 27B reported as negative result; model failed to converge beyond majority-class prediction.',
      italic=True, size=8, sa=12)

# --- Supervised ---
add_heading_n('Supervised deep learning performance', 2)

add_p(
    'Among supervised classifiers, Swin-T achieved the highest performance: accuracy 85.6% '
    '(95% CI 83.3-87.6%), sensitivity 85.2%, specificity 86.6%, F1\u2009=\u20090.896, '
    'MCC\u2009=\u20090.672, and AUC-ROC\u2009=\u20090.937 (Table 3). ResNet-50 achieved 76.0% accuracy '
    '(MCC\u2009=\u20090.527), while EfficientNet-B0 achieved 70.5% (MCC\u2009=\u20090.462). Notably, '
    'supervised models exhibited high specificity: even EfficientNet-B0 achieved 88.1%, in stark '
    'contrast to zero-shot VLMs. Training times were 5-6 minutes per model.'
)

# TABLE 3: Supervised
add_p('Table 3 | Supervised deep learning baseline performance.', bold=True, size=9, sa=4)
t3d = [
    ['Model', 'Accuracy', 'Sensitivity', 'Specificity', 'F1', 'MCC', 'AUC-ROC', 'Train time'],
    ['Swin-T', '0.856', '0.852', '0.866', '0.896', '0.672', '0.937', '5.3 min'],
    ['ResNet-50', '0.760', '0.720', '0.870', '0.814', '0.527', '0.889', '5.8 min'],
    ['EfficientNet-B0', '0.705', '0.640', '0.881', '0.760', '0.462', '0.849', '5.2 min'],
]
t3 = doc.add_table(rows=4, cols=8)
for i, rd in enumerate(t3d):
    for j, v in enumerate(rd):
        t3.rows[i].cells[j].text = v
fmt_table(t3)
for i in range(1, 4):
    t3.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
add_p('', sa=12)

# --- Cross-Tier ---
add_heading_n('Cross-tier comparison', 2)

# EDIT 1: Soften "statistically indistinguishable" → "no statistically significant difference"
add_p(
    'The comprehensive comparison across all three tiers is presented in Fig. 2 and Fig. 3. '
    'LLaVA-NeXT 7B achieved the numerically highest MCC (0.693) and accuracy (87.1%, 95% CI '
    '84.9-89.0%), narrowly exceeding Swin-T (MCC\u2009=\u20090.672; accuracy 85.6%, 95% CI '
    '83.3-87.6%). However, no statistically significant difference in accuracy was observed: '
    'the 95% confidence intervals overlap substantially, and a two-proportion z-test yields '
    'p\u2009=\u20090.33 for the 1.5 percentage-point gap. Performance was numerically similar across '
    'principal metrics, with near-identical AUC-ROC (0.938 vs. 0.937).'
)

add_p(
    'MCC reveals the true performance hierarchy on this imbalanced dataset (73.1% malignant). '
    'Raw accuracy inflates the apparent performance of models that predict predominantly malignant. '
    'MCC shows a clear separation: fine-tuned and supervised models (MCC 0.46-0.69) versus all '
    'zero-shot VLMs (MCC\u2009<\u20090.15). Training efficiency varied by orders of magnitude: supervised '
    'models trained in 5-6 minutes versus approximately 3 hours for VLM fine-tuning.'
)

# FIGURE 2
add_fig(f'{FIGS}/Fig2_MCC_comparison.png', 5.5,
    'Fig. 2 | Matthews correlation coefficient (MCC) across all model configurations. '
    'MCC accounts for all four confusion matrix quadrants and is robust to class imbalance. '
    'All zero-shot VLMs (blue) fall below 0.2 (near-random threshold), while fine-tuned VLMs (red) '
    'and supervised baselines (green) achieve MCC\u2009>\u20090.46. The Gemma 4 27B fine-tuning failure '
    '(grey) produced negative MCC.'
)

# FIGURE 3
add_fig(f'{FIGS}/Fig3_tier_comparison.png', 5.5,
    'Fig. 3 | Head-to-head comparison of the best model from each tier. '
    '(a) Accuracy, sensitivity, and specificity. The specificity gap between zero-shot (10.0%) '
    'and fine-tuned/supervised models (84-87%) is the defining finding. '
    '(b) MCC and AUC-ROC. Fine-tuned LLaVA-NeXT and supervised Swin-T achieve near-identical '
    'AUC-ROC (0.938 vs. 0.937). AUC is not available for zero-shot models (no probability outputs).'
)

# --- Prevalence ---
add_heading_n('Prevalence-adjusted clinical utility analysis', 2)

add_p(
    'Because the TN5000 test set has an atypically high malignant prevalence (73.1%), we performed '
    'a prevalence-adjusted modelling exercise to illustrate expected clinical utility at '
    'population-relevant prevalence levels (Table 4, Fig. 4). This analysis assumes that sensitivity '
    'and specificity transfer unchanged across populations, an assumption that may not hold in '
    'practice and should be interpreted as an illustrative exercise rather than a direct estimate '
    'of real-world performance.'
)

add_p(
    'Although the evaluated task corresponds to referred-nodule decision support rather than population '
    'screening, the 5% prevalence scenario illustrates how low specificity becomes particularly '
    'problematic in low-prevalence settings. At 5% prevalence, the best zero-shot VLM (Qwen3-VL-235B: '
    'sensitivity 98.2%, specificity 10.0%) would yield a PPV of only 5.4%, generating 17.4 false '
    'positives for every true malignancy detected. In contrast, the fine-tuned LLaVA-NeXT 7B achieves '
    'PPV of 22.9% at 5% prevalence, with 3.4 false positives per true positive, a five-fold reduction. '
    'At the more clinically representative 10-15% prevalence of a referred population, the fine-tuned '
    'model achieves PPV of 38.6-49.9% with FP/TP ratios of 1.0-1.6 (Table 4).'
)

# TABLE 4: Prevalence
add_p('Table 4 | Prevalence-adjusted clinical utility for selected models.', bold=True, size=9, sa=4)
t4d = [
    ['Model', 'Prevalence', 'PPV', 'NPV', 'FP per 100', 'FP/TP ratio'],
    ['Qwen3-VL-235B (best zero-shot)', '5%', '5.4%', '99.1%', '85.5', '17.4'],
    ['', '10%', '10.8%', '98.0%', '81.0', '8.2'],
    ['', '15%', '16.1%', '96.9%', '76.5', '5.2'],
    ['LLaVA-NeXT 7B (fine-tuned)', '5%', '22.9%', '99.3%', '14.8', '3.4'],
    ['', '10%', '38.6%', '98.5%', '14.0', '1.6'],
    ['', '15%', '49.9%', '97.6%', '13.3', '1.0'],
    ['Swin-T (supervised)', '5%', '25.1%', '99.1%', '12.7', '3.0'],
    ['', '10%', '41.4%', '98.1%', '12.1', '1.4'],
    ['', '15%', '52.9%', '97.1%', '11.4', '0.9'],
]
t4 = doc.add_table(rows=len(t4d), cols=6)
for i, rd in enumerate(t4d):
    for j, v in enumerate(rd):
        t4.rows[i].cells[j].text = v
fmt_table(t4)
for i in range(1, len(t4d)):
    t4.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
add_p(
    'FP per 100\u2009=\u2009expected false positives per 100 patients evaluated. PPV and NPV calculated '
    'using Bayes\u2019 theorem from observed sensitivity and specificity. This modelling assumes '
    'stable operating characteristics across populations.',
    italic=True, size=8, sa=12
)

# FIGURE 4
add_fig(f'{FIGS}/Fig4_prevalence_adjusted.png', 5.5,
    'Fig. 4 | Prevalence-adjusted clinical utility modelling. '
    '(a) Positive predictive value (PPV) as a function of malignancy prevalence. Zero-shot VLMs '
    '(blue) remain below 30% PPV even at 30% prevalence due to poor specificity. '
    '(b) False positives per true positive. At 5% prevalence, zero-shot VLMs generate ~18 false '
    'positives per true positive versus ~3 for fine-tuned and supervised models.'
)

# ============================================================
# DISCUSSION
# ============================================================
add_heading_n('Discussion', 1)

add_p(
    'The most clinically significant finding is the universal specificity failure of zero-shot VLMs '
    '(median 13.8%), which persists regardless of model family, size (7B to 235B parameters), or '
    'prompt engineering strategy. Several non-mutually-exclusive mechanisms may contribute. VLMs '
    'pre-trained on internet-scale corpora may encounter descriptions of thyroid nodules that '
    'disproportionately emphasise malignancy indicators and cautionary language; without domain '
    'calibration, models may inherit this textual prior. As Zhang et al.\u00b2\u00b2 have discussed, multimodal LLMs '
    'are substantially influenced by their pretraining data distributions. Alternatively, the bias may reflect '
    'a learned \u201csafe\u201d default: in medical contexts, predicting malignancy is the conservative '
    'choice that avoids the perceived cost of missed cancer. The full zero-shot results (Table 1) '
    'reveal that the effect of structured prompts (CoT, TI-RADS) on specificity was inconsistent '
    'and model-specific: in several model families, structured prompting reinforced rather than '
    'corrected the malignancy-prediction bias, while in others it had negligible or mixed effects.'
)

add_p(
    'Our results demonstrate that QLoRA fine-tuning can serve as an effective corrective mechanism '
    'for the specificity deficit in selected architectures. While every zero-shot VLM exhibited '
    'severe specificity collapse, QLoRA fine-tuning enabled LLaVA-NeXT 7B to achieve 84.4% specificity '
    'after training on only 3,500 images and modifying only 0.2% of model parameters (Fig. 5). '
    'LoRA\u2019s low-rank constraint likely provides implicit regularisation that prevents catastrophic '
    'forgetting while allowing task-specific decision boundaries. However, the failure of '
    'Gemma 4 27B cautions against assuming universal applicability.'
)

add_p(
    'We do not identify the specific cause of the Gemma 4 27B failure, but several '
    'architecture-specific factors may plausibly have contributed. Candidates include sensitivity '
    'to 4-bit NF4 quantisation at 27B scale (with possible distortion of the vision-language '
    'projection), mismatch between adapter placement and the model\u2019s image-language coupling '
    '(persistent majority-class collapse across multiple hyperparameter configurations is '
    'documented in Supplementary Appendix B.3), and the possibility that effective adaptation '
    'requires components, such as the vision encoder or cross-modal attention layers, that our '
    'LoRA targeting did not modify. The pattern is consistent with the broader interpretation '
    'that PEFT success depends on architecture-specific compatibility rather than parameter count.'
)

# FIGURE 5
add_fig(f'{FIGS}/Fig5_specificity_correction.png', 3.5,
    'Fig. 5 | Specificity correction through QLoRA fine-tuning. '
    'The zero-shot VLM median specificity (13.8%, dashed line) is corrected by fine-tuning: '
    'LLaVA-NeXT achieves 84.4% (+70.6 pp) and Qwen2.5-VL achieves 75.5% (+61.7 pp), '
    'approaching the supervised Swin-T baseline (86.6%).'
)

# EDIT 1 continued: Soften language throughout Discussion
add_p(
    'No statistically significant difference in accuracy was observed between fine-tuned LLaVA-NeXT '
    'and Swin-T (overlapping 95% CIs; p\u2009=\u20090.33), raising the question of what fine-tuned VLMs '
    'offer over simpler supervised classifiers. The supervised approach was substantially more '
    'practical: Swin-T trained in 5.3 minutes versus 179.7 minutes for LLaVA-NeXT, a 30-fold '
    'efficiency advantage; it required no quantisation infrastructure, and produced a smaller '
    'deployment footprint. Given numerically similar discriminative performance (no statistically '
    'significant difference in accuracy; overlapping 95% CIs; p\u2009=\u20090.33), a sharp question follows: '
    'for a single-task binary classification problem on cropped medical images, what is the '
    'clinical or technical justification for preferring a fine-tuned VLM over a supervised '
    'classifier? Fine-tuned VLMs retain a language generation architecture that could in principle '
    'support natural-language explanations, multi-task adaptation, or flexible instruction '
    'following. These are genuine potential advantages, but they are advantages we did not evaluate '
    'in this study, and it remains unknown whether task-specific QLoRA preserves them. Until '
    'VLM-specific capabilities are demonstrated to add clinical value on an identical dataset, '
    'conventional supervised models appear to be the more defensible choice for pure binary '
    'thyroid nodule classification on efficiency, reproducibility, and deployment grounds.'
)

add_p(
    'An important interpretive caveat: Tier 1 (zero-shot) and Tier 2 (fine-tuned) do not evaluate '
    'the same model families. The closed-source frontier VLMs evaluated zero-shot (Gemini, GPT, '
    'Claude, Grok, Llama 4 Scout) do not permit weight-level fine-tuning, and the open-weight models '
    'amenable to QLoRA (LLaVA-NeXT, Qwen2.5-VL, Gemma 4) were not uniformly available via the '
    'zero-shot API gateway. Tier 1 and Tier 2 should therefore not be interpreted as paired '
    'before-and-after comparisons within identical architectures, but rather as two separate '
    'characterisations of the current commercial and open-weight VLM ecosystems on this task.'
)

add_p(
    'Prompt engineering effects were inconsistent and model-specific (Table 1): TI-RADS guidance '
    'improved some models while severely degrading others (e.g., Claude Sonnet 4 dropped from 73.5% '
    'to 46.0%). Maximum accuracy gains from prompt optimisation were 3-5 percentage points, with '
    'non-negligible risk of degradation. In contrast, QLoRA fine-tuning produced 12-17 '
    'percentage-point accuracy gains.'
)

add_p(
    'The observation that structured TI-RADS prompting sometimes degraded rather than improved '
    'performance merits specific consideration. We advance this as hypothesis rather than explanation, '
    'but several mechanisms are plausible. First, TI-RADS-specific terminology (for example, '
    '\u201cpunctate echogenic foci\u201d, \u201ctaller-than-wide\u201d, \u201cspongiform\u201d) may be poorly grounded '
    'in the visual representations of general-purpose VLMs whose pretraining data contains limited '
    'exposure to curated thyroid sonographic lexicon. Second, a structured rubric may induce the '
    'model to hallucinate feature extraction, confidently assigning points for features the image '
    'does not display. Third, prompted stepwise reasoning may amplify, rather than calibrate, an '
    'existing malignancy-biased prior: each structured feature becomes an opportunity to accumulate '
    'suspicion, and the model converges toward a higher category. Fourth, the lexicon in TI-RADS '
    '(\u201csuspicious\u201d, \u201chighly suspicious\u201d) may itself shift the model\u2019s language-conditioned '
    'output distribution toward the malignant label. These hypotheses are not mutually exclusive and '
    'deserve targeted mechanistic study.'
)

add_p(
    'Our supervised baseline results are broadly consistent with the TN5000 publication\u00b2\u00b9. '
    'Recent work by Bahmane et al.\u2079 reported 89.7% accuracy on TN5000 using a hybrid '
    'EfficientNet-B3 model with GAN-based augmentation, though their evaluation used different '
    'train/test splits. Chen et al.\u00b2\u00b3 evaluated ChatGPT-4o and Claude 3-Opus on thyroid nodule '
    'classification from ultrasound images, reporting poor inter-rater agreement (kappa 0.034-0.116) '
    'and AUC below 57%, consistent with our zero-shot specificity collapse findings.'
)

add_p(
    'Although our findings are specific to thyroid ultrasound, the broad pattern, systematic '
    'specificity collapse in zero-shot VLMs coupled with selective recovery through PEFT, may '
    'plausibly extend to other narrowly defined medical imaging classification tasks where the '
    'relevant visual features are domain-specific and under-represented in internet-scale '
    'pretraining corpora. Likely candidates include dermatoscopic lesion classification, chest '
    'radiograph triage, and retinal disease grading. The extent of generalisation will depend on '
    'modality, label structure, clinical prevalence, and how well task-relevant visual features '
    'align with the model\u2019s pretraining distribution. We offer this extrapolation as a '
    'testable hypothesis rather than an established result; multi-task benchmarks across imaging '
    'modalities would be required to characterise the boundary conditions.'
)

# Limitations
add_heading_n('Limitations', 2)

add_p(
    'This study has several important limitations. First, all results are derived from a single '
    'public benchmark split from one institution, with no external validation. The conclusions are '
    'therefore benchmark-specific and should be interpreted as hypothesis-generating. Second, the '
    'TN5000 test set has high malignant prevalence (73.1%) that does not reflect clinical populations. '
    'Third, classification uses pre-cropped nodule images, bypassing the detection step required '
    'in clinical practice. Fourth, all fine-tuning experiments used a single random seed. Fifth, '
    'the Gemma 4 27B failure may be addressable through alternative PEFT strategies. Sixth, we did '
    'not evaluate natural-language explanation quality. Finally, computational requirements for VLM '
    'fine-tuning (A100-80GB, ~3 hours) remain substantial compared to supervised alternatives.'
)

# ============================================================
# METHODS (at end, Nature style)
# ============================================================
add_heading_n('Methods', 1)

add_heading_n('Study design and model accounting', 2)
add_p(
    'We conducted a systematic benchmark comprising three experimental tiers on a common dataset '
    'and test set. Table 5 provides complete model accounting.'
)

add_p('Table 5 | Complete model accounting.', bold=True, size=9, sa=4)
t5d = [
    ['Tier', 'Description', 'Unique models', 'Configurations'],
    ['1: Zero-shot VLMs', '12 VLMs \u00d7 3 prompts each', '12', '36'],
    ['2: Fine-tuned VLMs', 'QLoRA on 3 open-weight VLMs', '3', '3'],
    ['3: Supervised DL', 'ImageNet-pretrained classifiers', '3', '3'],
    ['Total', '', '18', '42'],
]
t5 = doc.add_table(rows=5, cols=4)
for i, rd in enumerate(t5d):
    for j, v in enumerate(rd):
        t5.rows[i].cells[j].text = v
fmt_table(t5)
for i in range(1, 5):
    t5.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t5.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
for j in range(4):
    for run in t5.rows[4].cells[j].paragraphs[0].runs:
        run.bold = True
add_p('', sa=6)

add_heading_n('Dataset', 2)
add_p(
    'We used the TN5000 dataset\u00b2\u00b9, a publicly available collection of 5,000 B-mode thyroid '
    'ultrasound images with biopsy-confirmed diagnoses from the National Cancer Center of China. '
    'We followed the authors\u2019 recommended partition: training (n\u2009=\u20093,500), validation '
    '(n\u2009=\u2009500), and test (n\u2009=\u20091,000). The test set comprises 731 malignant and 269 benign '
    'nodules (73.1% malignant prevalence). All images are pre-cropped greyscale B-mode ultrasound '
    'images; the detection step is not evaluated. The classification task corresponds to a '
    'decision-support scenario in which a nodule has already been identified and requires risk '
    'stratification.'
)

add_heading_n('Tier 1: zero-shot VLM evaluation', 2)
add_p(
    'Twelve VLMs spanning seven vendor ecosystems were evaluated: Gemini 2.0 Flash, Gemini 2.5 Flash, '
    'Gemini 2.5 Pro (Google); GPT-4.1, GPT-4.1 Mini, OpenAI o4-mini (OpenAI); Qwen2.5-VL-72B, '
    'Qwen3-VL-235B (Alibaba); Claude Sonnet 4, Claude Sonnet 4.6 (Anthropic); Grok 4 (xAI); and '
    'Llama 4 Scout (Meta). All models were accessed through the OpenRouter API gateway.'
)
add_p(
    'Each model was evaluated under three prompting conditions: (a) baseline: a direct '
    'classification request; (b) chain-of-thought (CoT): requiring systematic evaluation of five '
    'sonographic features before classification; and (c) TI-RADS guided: embedding the complete '
    'ACR TI-RADS point-based scoring criteria\u2074. Complete prompt templates are provided in '
    'Supplementary Appendix A.'
)

# EDIT 3: Restore zero-shot parsing pipeline details
add_p(
    'Model responses were parsed using a three-level extraction pipeline: '
    '(1) explicit \u201cmalignant\u201d/\u201cbenign\u201d keyword matching (case-insensitive); '
    '(2) TI-RADS category extraction (TR4/TR5 mapped to malignant, TR1/TR2/TR3 to benign); and '
    '(3) rule-based sentiment classification using predefined term lists (see Supplementary Appendix C '
    'for the complete rule set). Level 3 was required for fewer than 0.3% of total classifications. '
    'Responses that could not be classified after all stages were recorded as parsing failures '
    '(<1% for all models). Refusal rates were negligible (<0.5%).',
    sa=6
)

add_heading_n('Tier 2: QLoRA fine-tuning of VLMs', 2)
add_p(
    'Three open-weight VLMs were selected: LLaVA-NeXT 7B (Mistral backbone), '
    'Qwen2.5-VL-7B, and Gemma 4 27B. Selection was constrained by publicly available '
    'weights compatible with Hugging Face PEFT infrastructure. The closed-source VLMs in Tier 1 do '
    'not permit weight-level fine-tuning. Consequently, zero-shot and fine-tuned tiers do not compare '
    'the same model families. Zero-shot LLaVA-NeXT 7B was not included in the zero-shot tier because '
    'it was not available through the OpenRouter API at the time of evaluation.'
)

add_p(
    'Fine-tuning used QLoRA: 4-bit NormalFloat (NF4) quantisation via bitsandbytes '
    'with LoRA. All models were fine-tuned for 3 epochs on the training set (n\u2009=\u20093,500) '
    'using SFTTrainer from the TRL library. Hyperparameters: LoRA rank r\u2009=\u200916, '
    'alpha\u2009=\u200932, effective batch size\u2009=\u200916, learning rate\u2009=\u20092\u2009\u00d7\u200910\u207b\u2074 with '
    'cosine scheduling, maximum sequence length 2,048 tokens. LoRA modules were applied to all linear '
    'layers in the language model backbone; vision encoder weights were preserved. A '
    'WeightedRandomSampler addressed class imbalance (71.4% malignant).'
)

# EDIT 4: Remove augmentation from Tier 2 (it was supervised-only)
add_p(
    'No image-level data augmentation was applied during VLM fine-tuning; augmentation was used only '
    'in the supervised training pipeline (Tier 3). The final checkpoint after 3 epochs was used for '
    'evaluation. Inference used greedy decoding (temperature\u2009=\u20090). Binary predictions were '
    'extracted via keyword matching; parsing error rates were 0% for both LLaVA-NeXT and Qwen2.5-VL. '
    'All fine-tuning was performed on NVIDIA A100-80GB GPUs via Modal serverless infrastructure, '
    'with training times of 161-180 minutes per model. The random seed was fixed at 42.',
    sa=6
)

add_heading_n('Tier 3: supervised deep learning baselines', 2)
add_p(
    'Three architectures were trained: ResNet-50, EfficientNet-B0, and Swin-T, '
    'following established deep learning\u00b2\u2074 and transfer learning\u00b2\u2075 methodology for medical '
    'imaging, including vision transformer approaches for ultrasound\u00b2\u2076. '
    'All were initialised from ImageNet-pretrained weights. A two-phase strategy was used: '
    '(1) frozen backbone for 5 epochs, then (2) full fine-tuning at one-tenth the initial learning '
    'rate. Training used cross-entropy loss with class weights inversely proportional to frequency, '
    'combined with WeightedRandomSampler. Hyperparameters: 30 maximum epochs, batch size 32, '
    'learning rate 1\u2009\u00d7\u200910\u207b\u2074, weight decay 1\u2009\u00d7\u200910\u207b\u2074, early stopping with '
    'patience of 7 based on validation loss. Data augmentation: random horizontal flip (p\u2009=\u20090.5), '
    'rotation (\u00b115\u00b0), colour jitter (brightness, contrast, saturation, hue each \u00b10.2). '
    'Images resized to 224\u2009\u00d7\u2009224 pixels. Seed fixed at 42. Training times: 5-6 minutes per model.',
    sa=6
)

add_heading_n('Statistical analysis', 2)
add_p(
    'For each model, we computed accuracy, sensitivity, specificity, precision, F1, Matthews '
    'correlation coefficient (MCC), and AUC-ROC. The 95% Wilson score confidence intervals were '
    'calculated for proportion-based metrics\u00b2\u2077. MCC was selected as a primary metric because it '
    'accounts for all four confusion matrix quadrants and is robust to class imbalance\u00b2\u2078. '
    'For head-to-head comparisons between the best-performing fine-tuned and supervised models, '
    'formal hypothesis testing was used: the two-proportion z-test for accuracy at '
    '\u03b1\u2009=\u20090.05. Confidence intervals are reported to convey the precision of individual '
    'point estimates. Prevalence-adjusted utility was modelled using Bayes\u2019 theorem at three '
    'prevalence levels: 5%, 10%, and 15%.'
)

# ============================================================
# CONCLUSION
# ============================================================
add_heading_n('Conclusion', 1)

add_p(
    'This three-tier benchmark of 18 models across 42 configurations on the TN5000 dataset '
    'establishes three findings. First, on this benchmark, current zero-shot VLMs did not reliably '
    'classify thyroid nodules: all 12 models exhibited severe specificity collapse (median 13.8%). Second, QLoRA '
    'fine-tuning corrected this deficit convincingly in one architecture (LLaVA-NeXT 7B: '
    'MCC\u2009=\u20090.693, AUC\u2009=\u20090.938) and partially in another (Qwen2.5-VL-7B: MCC\u2009=\u20090.539), '
    'while failing entirely in a third (Gemma 4 27B). Third, model architecture and PEFT '
    'compatibility, rather than parameter count, may influence fine-tuning success. For '
    'single-task binary classification on cropped thyroid ultrasound images, conventional '
    'supervised models may remain the more practical choice unless VLM-specific capabilities '
    '(explainability, multi-task adaptation, natural-language interaction) are shown to add '
    'clinical value on identical datasets.'
)

add_p(
    'These findings require confirmation through multi-centre external validation. Key priorities: '
    'prospective evaluation on independent datasets with representative prevalence; assessment of '
    'whether fine-tuned VLMs retain useful language capabilities; and systematic investigation of '
    'which architectural features predict successful domain adaptation via PEFT.'
)

# ============================================================
# REQUIRED npj SECTIONS
# ============================================================
add_heading_n('Data availability', 1)
add_p(
    'To ensure the fairness and reproducibility of the comparison, this research did not involve '
    'the collection of new patient data. The TN5000 thyroid ultrasound dataset utilised in this '
    'study is publicly available and was accessed under its original licence as described in '
    'Zhang et al.\u00b2\u00b9 (https://doi.org/10.6084/m9.figshare.27962824). The dataset was used '
    'throughout the experiments in strict adherence to its data use agreement. The performance of '
    'closed-source vision-language models (Gemini, GPT, Claude, Grok, Llama 4 Scout, Qwen2.5-VL-72B, '
    'Qwen3-VL-235B) was obtained through the OpenRouter API gateway. All open-weight models '
    '(LLaVA-NeXT 7B, Qwen2.5-VL-7B, Gemma 4 27B) and supervised deep learning baselines (ResNet-50, '
    'EfficientNet-B0, Swin-T) were deployed locally on NVIDIA A100-80GB GPUs via Modal serverless '
    'infrastructure. Processed model outputs, prediction files, and evaluation logs supporting the '
    'findings of this study are available in the GitHub repository listed under Code Availability.'
)

add_heading_n('Code availability', 1)
add_p(
    'The code supporting this benchmarking study, including all fine-tuning scripts, model '
    'configurations, inference pipelines, prompt templates, output parsing logic, evaluation '
    'scripts, statistical analyses, and figure-generation code, is available online at '
    '[GitHub repository URL to be inserted prior to publication]. Up-to-date benchmark results '
    'are additionally provided on the project page.'
)

add_heading_n('Author contributions', 1)
add_p(
    'F.W.A. led the study design, manuscript conceptualisation, conducted the analyses, '
    'and wrote the main manuscript text. F.A. contributed substantially to writing the main '
    'manuscript, the analytic framework and its formulation, and theoretical conceptualisation. '
    'F.W.A. and F.A. contributed to editing of the manuscript prior to publication. '
    'B.M. reviewed and edited the manuscript prior to submission.'
)

add_heading_n('Competing interests', 1)
add_p('The authors declare no competing interests.')

add_heading_n('Ethics declarations', 1)
add_p(
    'This study used the publicly available TN5000 dataset, which was collected and de-identified '
    'by the original authors. No additional ethics approval was required for this secondary analysis '
    'of publicly available, de-identified data.'
)

add_heading_n('Acknowledgements', 1)
add_p(
    'The authors acknowledge the developers of the TN5000 dataset for making this publicly '
    'available benchmark possible. Computational resources for parameter-efficient fine-tuning '
    'and supervised training were provided through Modal serverless GPU infrastructure.'
)

# ============================================================
# REFERENCES (Nature numbered style)
# ============================================================
add_heading_n('References', 1)

refs = [
    # 1-9: unchanged
    'Durante, C. et al. The Diagnosis and Management of Thyroid Nodules: A Review. JAMA 319, 914-924 (2018). PMID: 29509871.',
    'Pizzato, M. et al. The epidemiological landscape of thyroid cancer worldwide: GLOBOCAN estimates for incidence and mortality rates in 2020. Lancet Diabetes Endocrinol. 10, 264-272 (2022). PMID: 35271818.',
    'Cibas, E. S. & Ali, S. Z. The 2017 Bethesda System for Reporting Thyroid Cytopathology. Thyroid 27, 1341-1346 (2017). PMID: 29091573.',
    'Tessler, F. N. et al. ACR Thyroid Imaging, Reporting and Data System (TI-RADS): White Paper of the ACR TI-RADS Committee. J. Am. Coll. Radiol. 14, 587-595 (2017). PMID: 28372962.',
    'Bai, Z. et al. Thyroid nodules risk stratification through deep learning based on ultrasound images. Med. Phys. 47, 6355-6365 (2020). PMID: 33089513.',
    'Hung, W. C. et al. Thyroid Nodule Detection and Classification on Small Datasets: An Ensemble Deep Learning Approach with Attention Mechanism and Focal Loss. Diagnostics 16, 825 (2026). PMID: 41897558.',
    'Liu, X. et al. A dual-branch deep learning framework with Mask-Guided Attention for thyroid nodule classification in ultrasound images. Front. Med. 13, 1694174 (2026). PMID: 41852538.',
    'Kim, J. et al. Artificial intelligence-assisted risk stratification of thyroid nodules with atypia of undetermined significance. Eur. Thyroid J. 15, ETJ-25-0268 (2026). PMID: 41553093.',
    'Bahmane, K., Bhattacharya, S. & Chaouki, A. B. Evaluation of a Hybrid CNN Model for Automatic Detection of Malignant and Benign Lesions. Medicina 61, 2036 (2025). PMID: 41303872.',
    # 10-12: unchanged
    'Jiao, L. et al. Foundation Models Meet Medical Image Interpretation. Research 9, 1024 (2026). PMID: 41767596.',
    'Shamshad, F. et al. Transformers in medical imaging: A survey. Med. Image Anal. 88, 102802 (2023). PMID: 37315483.',
    'Aljaddouh, B., Malathi, D. & Alaswad, F. Multimodal vision-language models in chest x-ray analysis: a study of generalization, supervision, and robustness. Biomed. Eng. Lett. 16, 517-537 (2026). PMID: 41890265.',
    # 13: was 18 (Meddeb) - VLMs neuroradiology
    'Meddeb, A. et al. Evaluating the diagnostic accuracy of vision language models for neuroradiological image interpretation. npj Digit. Med. 8, 666 (2025). PMID: 41249440.',
    # 14: was 13 (Boghosian) - VLMs dermatology
    'Boghosian, T. et al. Diagnostic Accuracy and Pitfalls of Publicly Available Artificial Intelligence Models for Nail Disorders. J. Drugs Dermatol. 25, 349-356 (2026). PMID: 41931690.',
    # 15: was 14 (Zhao) - pediatric imaging
    'Zhao, Y. et al. Evidence-Guided Diagnostic Reasoning for Pediatric Chest Radiology Based on Multimodal Large Language Models. J. Imaging 12, 111 (2026). PMID: 41892913.',
    # 16: was 19 (Ko) - VLMs orthopedic
    'Ko, S. et al. Benchmarking open-source vision language models in orthopedic in-training examination. Clin. Orthop. Surg. 18, 159-166 (2025). PMID: 41647505.',
    # 17: was 20 (Liu) - LLM dentistry
    'Liu, X. et al. Developing and evaluating multimodal large language model for orthopantomography analysis to support clinical dentistry. Cell Rep. Med. 7, 102652 (2026). PMID: 41850234.',
    # 18: was 15 (Hou) - LLaMA fine-tuning
    'Hou, Y. et al. Fine-tuning a local LLaMA-3 large language model for automated privacy-preserving physician letter generation in radiation oncology. Front. Artif. Intell. 7, 1493716 (2025). PMID: 39877751.',
    # 19: was 16 (Sahoo) - LLMs biomedicine
    'Sahoo, S. S. et al. Large language models for biomedicine: foundations, opportunities, challenges, and best practices. J. Am. Med. Inform. Assoc. 31, 2114-2124 (2024). PMID: 38657567.',
    # 20: was 24 (Liu) - ChatOCT / LLM fine-tuning
    'Liu, C. et al. ChatOCT: embedded clinical decision support systems for optical coherence tomography in offline and resource-limited settings. J. Med. Syst. 49, 59 (2025). PMID: 40332685.',
    # 21: was 17 (Zhang) - TN5000 dataset
    'Zhang, H. et al. TN5000: An Ultrasound Image Dataset for Thyroid Nodule Detection and Classification. Sci. Data 12, 1437 (2025). PMID: 40819171.',
    # 22: was 27 (Zhang) - multimodal LLMs bioimage
    'Zhang, S. et al. Multimodal large language models for bioimage analysis. Nat. Methods 21, 1390-1393 (2024). PMID: 39122942.',
    # 23: was 28 (Chen) - ChatGPT thyroid
    'Chen, Z. et al. Assessing the feasibility of ChatGPT-4o and Claude 3-Opus in thyroid nodule classification based on ultrasound images. Endocrine 87, 1041-1049 (2025). PMID: 39394537.',
    # 24: was 21 (Litjens) - DL survey
    'Litjens, G. et al. A survey on deep learning in medical image analysis. Med. Image Anal. 42, 60-88 (2017). PMID: 28778026.',
    # 25: was 22 (Kim) - transfer learning
    'Kim, H. E. et al. Transfer learning for medical image classification: a literature review. BMC Med. Imaging 22, 69 (2022). PMID: 35418051.',
    # 26: was 23 (Vafaeezadeh) - US vision transformers
    'Vafaeezadeh, M. et al. Ultrasound image analysis with vision transformers - review. Diagnostics 14, 542 (2024). PMID: 38473014.',
    # 27: was 25 (Wilson)
    'Wilson, E. B. Probable inference, the law of succession, and statistical inference. J. Am. Stat. Assoc. 22, 209-212 (1927).',
    # 28: was 26 (Chicco) - MCC
    'Chicco, D. & Jurman, G. The advantages of the Matthews correlation coefficient (MCC) over F1 score and accuracy in binary classification evaluation. BMC Genomics 21, 6 (2020). PMID: 31898477.',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    nr = p.add_run(f'{i}. '); nr.font.name = 'Arial'; nr.font.size = Pt(8)
    rr = p.add_run(ref); rr.font.name = 'Arial'; rr.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)


# ============================================================
# PAGE BREAK → SUPPLEMENTARY INFORMATION
# ============================================================
doc.add_page_break()

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run('Supplementary Information')
sr.font.name = 'Arial'; sr.font.size = Pt(16); sr.bold = True
sp.paragraph_format.space_after = Pt(6)

add_p(
    'Parameter-Efficient Fine-Tuning Corrects Specificity Collapse in '
    'Zero-Shot Vision-Language Models for Thyroid Ultrasound Classification',
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, sa=24
)

# ---- Supplementary Appendix A: Prompt Templates ----
add_heading_n('Supplementary Appendix A: Prompt Templates', 1)

add_p('Supplementary Table S1 | Prompt templates used for zero-shot VLM evaluation.', bold=True, size=9, sa=6)

add_p('A.1 Baseline Prompt', bold=True, sa=2)
add_p(
    'You are a medical imaging expert. Please examine this thyroid ultrasound image carefully. '
    'Classify the thyroid nodule shown as either "malignant" or "benign". '
    'Provide your classification as a single word: malignant or benign.',
    size=9, sa=8
)

add_p('A.2 Chain-of-Thought (CoT) Prompt', bold=True, sa=2)
add_p(
    'You are a medical imaging expert. Please examine this thyroid ultrasound image systematically. '
    'Evaluate the following five sonographic features step by step:\n'
    '1. Composition (solid, mixed cystic-solid, cystic, spongiform)\n'
    '2. Echogenicity (hyperechoic, isoechoic, hypoechoic, very hypoechoic)\n'
    '3. Shape (wider-than-tall, taller-than-wide)\n'
    '4. Margin (smooth, ill-defined, lobulated, irregular, extrathyroidal extension)\n'
    '5. Echogenic foci (none, macrocalcifications, peripheral calcifications, punctate echogenic foci)\n\n'
    'After evaluating each feature, provide your final classification as: malignant or benign.',
    size=9, sa=8
)

add_p('A.3 TI-RADS Guided Prompt', bold=True, sa=2)
add_p(
    'You are a radiologist using ACR TI-RADS to classify thyroid nodules. '
    'Examine this ultrasound image and assign points using the ACR TI-RADS scoring system:\n'
    '\u2022 Composition: Cystic/spongiform (0), Mixed cystic-solid (1), Solid/almost solid (2)\n'
    '\u2022 Echogenicity: Anechoic (0), Hyper/isoechoic (1), Hypoechoic (2), Very hypoechoic (3)\n'
    '\u2022 Shape: Wider-than-tall (0), Taller-than-wide (3)\n'
    '\u2022 Margin: Smooth (0), Ill-defined (0), Lobulated/irregular (2), Extrathyroidal (3)\n'
    '\u2022 Echogenic foci: None/large comet-tail (0), Macrocalcifications (1), Peripheral (2), Punctate (3)\n\n'
    'Sum the points and assign a TI-RADS category:\n'
    'TR1 (0 pts): Benign | TR2 (2 pts): Not suspicious | TR3 (3 pts): Mildly suspicious\n'
    'TR4 (4-6 pts): Moderately suspicious | TR5 (7+ pts): Highly suspicious\n\n'
    'Based on the TI-RADS category, classify as: malignant or benign.',
    size=9, sa=12
)

# ---- Supplementary Appendix B: Model-Specific Fine-Tuning Details ----
add_heading_n('Supplementary Appendix B: Model-Specific Fine-Tuning Adaptations', 1)

add_p('B.1 LLaVA-NeXT 7B (Mistral backbone)', bold=True, sa=2)
add_p(
    'The Mistral chat template was applied manually in the data collator rather than through the '
    'processor\u2019s template function, to ensure correct token formatting. The conversation format '
    'used [INST] and [/INST] delimiters as specified by the Mistral architecture. Vision tokens '
    'were injected at the position indicated by the <image> placeholder in the user message. '
    'Training completed in 179.7 minutes (3 epochs, 656 steps). Final checkpoint was used for evaluation.',
    size=9, sa=8
)

add_p('B.2 Qwen2.5-VL-7B', bold=True, sa=2)
add_p(
    'Truncation was disabled for vision tokens (approximately 1,280 tokens per image) to prevent '
    'image-text token mismatches, requiring a minimum sequence length of 2,048. The Qwen2.5-VL '
    'processor automatically handles image resizing and patch extraction. Training completed in '
    '161.7 minutes (3 epochs, 656 steps). Final checkpoint was used for evaluation.',
    size=9, sa=8
)

add_p('B.3 Gemma 4 27B', bold=True, sa=2)
add_p(
    'Gemma 4 27B required post-hoc removal of LoRA layers inadvertently applied to the vision '
    'encoder\u2019s patch embedding projection, which received uint8 pixel values incompatible with '
    'the dropout operations in LoRA layers. Multiple hyperparameter configurations were attempted: '
    'learning rates of 2\u2009\u00d7\u200910\u207b\u2074 and 2\u2009\u00d7\u200910\u207b\u2075, LoRA ranks 16 and 32, '
    'and 3-5 training epochs. All configurations produced predictions collapsing toward the '
    'majority (malignant) class with specificity <15%. The model\u2019s fine-tuning failure is '
    'reported as a negative result to document that model scale does not guarantee PEFT success.',
    size=9, sa=12
)

# ---- Supplementary Appendix C: Output Parsing Pipeline ----
add_heading_n('Supplementary Appendix C: Zero-Shot Output Parsing Pipeline', 1)

add_p(
    'Model responses were parsed using a hierarchical extraction pipeline with three levels:',
    size=9, sa=4
)

add_p(
    'Level 1: Keyword matching: Responses were searched (case-insensitive) for the keywords '
    '\u201cmalignant\u201d and \u201cbenign\u201d. If exactly one keyword was found, the corresponding label was '
    'assigned. If both keywords appeared, the final occurrence was used (as models typically state '
    'their conclusion at the end of the response).',
    size=9, sa=4
)

add_p(
    'Level 2: TI-RADS category extraction: For responses mentioning TI-RADS categories without '
    'explicit malignant/benign keywords, categories TR4 and TR5 were mapped to \u201cmalignant\u201d, '
    'and TR1, TR2, and TR3 were mapped to \u201cbenign\u201d. This mapping aligns with the ACR TI-RADS '
    'recommendation thresholds for biopsy consideration.',
    size=9, sa=4
)

add_p(
    'Level 3: Rule-based sentiment classification: For the small number of responses that could '
    'not be parsed at Levels 1 or 2 (typically verbose narrative responses without explicit keywords '
    'or TI-RADS categories), a predefined rule-based classifier was applied. Responses containing '
    'terms from a fixed concerning-language list (\u201csuspicious\u201d, \u201cconcerning\u201d, '
    '\u201cirregular\u201d, \u201chypoechoic\u201d, \u201ctaller-than-wide\u201d, \u201ccalcification\u201d, '
    '\u201chighly suspicious\u201d, \u201crecommend biopsy\u201d) were mapped to \u201cmalignant\u201d; '
    'responses containing reassuring terms (\u201cnormal\u201d, \u201cbenign-appearing\u201d, '
    '\u201cno concerning features\u201d, \u201cisoechoic\u201d, \u201ccystic\u201d, \u201cspongiform\u201d, '
    '\u201cno further workup\u201d) were mapped to \u201cbenign\u201d. This rule set was defined '
    'prior to evaluation and applied identically to all models. Level 3 parsing was required for '
    'fewer than 0.3% of total classifications (approximately 10-15 responses per model across '
    'all prompts); its impact on aggregate metrics is therefore negligible.',
    size=9, sa=4
)

add_p(
    'Unparseable responses: Responses that could not be classified after all three levels were '
    'recorded as parsing failures and excluded from metric calculations. Parsing failure rates '
    'were below 1% for all models. Model refusals (responses explicitly declining to classify, '
    'such as \u201cI cannot provide medical diagnoses\u201d) were recorded separately and were '
    'negligible (<0.5%) across all evaluated models.',
    size=9, sa=12
)

# ---- Supplementary Appendix D: Confusion Matrices ----
add_heading_n('Supplementary Appendix D: Confusion Matrices', 1)

add_p(
    'Supplementary Figure S1 presents confusion matrices for the top-performing models in each tier.',
    size=9, sa=6
)

# Check for existing confusion matrix images
cm_figs = [
    ('/Users/fahadwali/modal_results/TN5000_all_results/llava_confusion_matrix.png',
     'Supplementary Fig. S1a | Confusion matrix for LLaVA-NeXT 7B (QLoRA fine-tuned). '
     'TP=644, TN=227, FP=42, FN=87. n=1,000.'),
    ('/Users/fahadwali/modal_results/TN5000_all_results/qwen25vl_confusion_matrix.png',
     'Supplementary Fig. S1b | Confusion matrix for Qwen2.5-VL-7B (QLoRA fine-tuned). '
     'TP=599, TN=203, FP=66, FN=132. n=1,000.'),
    ('/Users/fahadwali/modal_results/TN5000_all_results/confusion_matrix_swin_t.png',
     'Supplementary Fig. S1c | Confusion matrix for Swin-T (supervised). '
     'TP=623, TN=233, FP=36, FN=108. n=1,000.'),
]

for path, caption in cm_figs:
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(3.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_p(caption, italic=True, size=8, sa=10)

# ---- Supplementary Appendix E: Training Curves ----
add_heading_n('Supplementary Appendix E: Training Curves', 1)

tc_figs = [
    ('/Users/fahadwali/modal_results/TN5000_all_results/llava_training_curves.png',
     'Supplementary Fig. S2a | Training loss curves for LLaVA-NeXT 7B QLoRA fine-tuning (3 epochs, 179.7 min).'),
    ('/Users/fahadwali/modal_results/TN5000_all_results/qwen25vl_training_curves.png',
     'Supplementary Fig. S2b | Training loss curves for Qwen2.5-VL-7B QLoRA fine-tuning (3 epochs, 161.7 min).'),
    ('/Users/fahadwali/modal_results/TN5000_all_results/training_curves_swin_t.png',
     'Supplementary Fig. S2c | Training and validation loss curves for Swin-T supervised training '
     '(best epoch 16/23, early stopping triggered).'),
    ('/Users/fahadwali/modal_results/TN5000_all_results/training_curves_resnet50.png',
     'Supplementary Fig. S2d | Training and validation loss curves for ResNet-50 (best epoch 16/23).'),
    ('/Users/fahadwali/modal_results/TN5000_all_results/training_curves_efficientnet_b0.png',
     'Supplementary Fig. S2e | Training and validation loss curves for EfficientNet-B0 (30 epochs, no early stopping).'),
]

for path, caption in tc_figs:
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(4.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_p(caption, italic=True, size=8, sa=10)

# ---- Supplementary Appendix F: Hyperparameter Summary ----
add_heading_n('Supplementary Appendix F: Hyperparameter Summary', 1)

add_p('Supplementary Table S2 | QLoRA fine-tuning hyperparameters.', bold=True, size=9, sa=4)
hp_data = [
    ['Parameter', 'LLaVA-NeXT 7B', 'Qwen2.5-VL-7B', 'Gemma 4 27B'],
    ['Base model', 'llava-hf/llava-v1.6-mistral-7b-hf', 'Qwen/Qwen2.5-VL-7B-Instruct', 'google/gemma-4-27b-it'],
    ['Quantisation', '4-bit NF4', '4-bit NF4', '4-bit NF4'],
    ['LoRA rank (r)', '16', '16', '16 / 32'],
    ['LoRA alpha', '32', '32', '32 / 64'],
    ['LoRA dropout', '0.05', '0.05', '0.05'],
    ['LoRA targets', 'All linear layers', 'All linear layers', 'All linear layers'],
    ['Learning rate', '2e-4', '2e-4', '2e-4 / 2e-5'],
    ['LR scheduler', 'Cosine', 'Cosine', 'Cosine'],
    ['Epochs', '3', '3', '3 / 5'],
    ['Effective batch size', '16', '16', '16'],
    ['Max seq length', '2,048', '2,048', '2,048'],
    ['Gradient checkpointing', 'Yes', 'Yes', 'Yes'],
    ['Training time', '179.7 min', '161.7 min', '~300 min (total)'],
    ['GPU', 'A100-80GB', 'A100-80GB', 'A100-80GB'],
    ['Seed', '42', '42', '42'],
    ['Converged', 'Yes', 'Yes', 'No'],
]
ts2 = doc.add_table(rows=len(hp_data), cols=4)
for i, rd in enumerate(hp_data):
    for j, v in enumerate(rd):
        ts2.rows[i].cells[j].text = v
fmt_table(ts2, hdr_color='2F5496')
for i in range(1, len(hp_data)):
    ts2.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
add_p('', sa=8)

add_p('Supplementary Table S3 | Supervised training hyperparameters.', bold=True, size=9, sa=4)
sup_hp = [
    ['Parameter', 'Value'],
    ['Architectures', 'ResNet-50, EfficientNet-B0, Swin-T'],
    ['Pre-training', 'ImageNet-1K'],
    ['Phase 1 (frozen)', '5 epochs, classification head only'],
    ['Phase 2 (full)', 'Up to 25 additional epochs at LR/10'],
    ['Initial learning rate', '1e-4'],
    ['Optimiser', 'Adam'],
    ['Weight decay', '1e-4'],
    ['Batch size', '32'],
    ['Loss function', 'Cross-entropy with class weights'],
    ['Class balancing', 'WeightedRandomSampler'],
    ['Early stopping', 'Patience = 7 (validation loss)'],
    ['Image size', '224 \u00d7 224'],
    ['Augmentation', 'HFlip (p=0.5), Rotation (\u00b115\u00b0), Colour jitter (\u00b10.2)'],
    ['Seed', '42'],
]
ts3 = doc.add_table(rows=len(sup_hp), cols=2)
for i, rd in enumerate(sup_hp):
    for j, v in enumerate(rd):
        ts3.rows[i].cells[j].text = v
fmt_table(ts3, hdr_color='2F5496')
for i in range(1, len(sup_hp)):
    ts3.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    ts3.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
add_p('', sa=12)

# ---- Supplementary Appendix G: Prevalence Calculations ----
add_heading_n('Supplementary Appendix G: Prevalence-Adjusted Calculations', 1)

add_p(
    'PPV and NPV were calculated using Bayes\u2019 theorem:',
    size=9, sa=4
)
add_p(
    'PPV = (Sensitivity \u00d7 Prevalence) / '
    '(Sensitivity \u00d7 Prevalence + (1 \u2212 Specificity) \u00d7 (1 \u2212 Prevalence))',
    size=9, sa=2
)
add_p(
    'NPV = (Specificity \u00d7 (1 \u2212 Prevalence)) / '
    '(Specificity \u00d7 (1 \u2212 Prevalence) + (1 \u2212 Sensitivity) \u00d7 Prevalence)',
    size=9, sa=2
)
add_p(
    'FP per 100 = (1 \u2212 Specificity) \u00d7 (1 \u2212 Prevalence) \u00d7 100',
    size=9, sa=2
)
add_p(
    'FP/TP ratio = ((1 \u2212 Specificity) \u00d7 (1 \u2212 Prevalence)) / (Sensitivity \u00d7 Prevalence)',
    size=9, sa=8
)

add_p(
    'These calculations assume that sensitivity and specificity measured on the TN5000 test set '
    '(73.1% malignant prevalence) remain stable when applied to populations with different prevalence. '
    'This assumption may not hold in practice due to spectrum bias and case-mix effects, and the '
    'results should be interpreted as illustrative modelling rather than predictive estimates of '
    'real-world performance.',
    size=9, sa=12
)

# ============================================================
# SAVE
# ============================================================
outpath = '/Users/fahadwali/Downloads/TN5000/TN5000_VLM_Benchmark_npjDigMed.docx'
doc.save(outpath)
print(f'\n\u2705 Final manuscript saved: {outpath}')
print(f'   Main text: ~5,000 words')
print(f'   Tables: 5 main + 3 supplementary')
print(f'   Figures: 5 main + supplementary (confusion matrices, training curves)')
print(f'   Supplementary appendices: A-G')
print(f'   References: {len(refs)}')
print(f'   Format: npj Digital Medicine')
