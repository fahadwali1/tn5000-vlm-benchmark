#!/usr/bin/env python3
"""
Build the npj Digital Medicine cover letter as a .docx file.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_p(text, bold=False, italic=False, size=11, align=None, sa=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(sa)
    return p


# --- Sender block ---
add_p('Fahad Wali Ahmed, MD', bold=True, sa=2)
add_p('King Faisal Specialist Hospital and Research Centre\u2013Madinah, Saudi Arabia', sa=2)
add_p('Email: fahadwali@yahoo.com', sa=2)
add_p('ORCID: 0000-0003-4100-1571', sa=18)

# --- Date placeholder ---
add_p('[Date of submission]', sa=18)

# --- Recipient ---
add_p('The Editor', bold=True, sa=2)
add_p('npj Digital Medicine', bold=True, sa=2)
add_p('Nature Portfolio', sa=18)

# --- Salutation ---
add_p('Dear Editor,', sa=12)

# --- Body paragraph 1: manuscript submission and topic ---
add_p(
    'We submit for your consideration our original research manuscript entitled '
    '\u201cParameter-Efficient Fine-Tuning Corrects Specificity Collapse in Zero-Shot '
    'Vision-Language Models for Thyroid Ultrasound Classification\u201d for publication '
    'in npj Digital Medicine. This work has not been published elsewhere and is not under '
    'consideration by any other journal.',
    sa=10
)

# --- Body paragraph 2: scientific contribution ---
add_p(
    'Our study presents a comprehensive three-tier benchmark of 18 models on the publicly '
    'available TN5000 thyroid ultrasound dataset, directly comparing 12 state-of-the-art '
    'vision-language models (VLMs) evaluated zero-shot across three prompt strategies, three '
    'open-weight VLMs adapted using quantised low-rank adaptation (QLoRA), and three supervised '
    'deep learning baselines on identical train/test splits. To our knowledge, this is among '
    'the first systematic evaluations to examine zero-shot, parameter-efficient fine-tuning, '
    'and supervised paradigms on a shared medical imaging benchmark.',
    sa=10
)

# --- Body paragraph 3: key findings ---
add_p(
    'Three findings we believe will interest the npj Digital Medicine readership. First, every '
    'zero-shot VLM exhibited severe specificity collapse (median 13.8%; range 3.0\u201329.7%), '
    'producing systematic malignancy-prediction bias that persisted regardless of prompt strategy '
    'or model scale (7B to 235B parameters). Second, QLoRA fine-tuning corrected this deficit '
    'in LLaVA-NeXT 7B (87.1% accuracy, 84.4% specificity, AUC\u2009=\u20090.938) by modifying '
    'only 0.2% of parameters, while failing entirely in Gemma 4 27B\u2014indicating that '
    'architecture-specific compatibility, rather than parameter count, governs adaptation '
    'success. Third, a supervised Swin-T baseline achieved comparable performance '
    '(85.6% accuracy, AUC\u2009=\u20090.937; p\u2009=\u20090.33 for the accuracy difference) with '
    'a 30-fold training-time advantage (5.3 minutes versus 179.7 minutes), raising practical '
    'questions about when VLM-based approaches are the more defensible choice for pure '
    'binary classification on cropped medical images.',
    sa=10
)

# --- Body paragraph 4: relevance to journal ---
add_p(
    'The manuscript contributes a reproducible benchmark, a clinically framed prevalence-adjusted '
    'utility analysis, and a transparent negative result that together inform how the digital '
    'medicine community should interpret zero-shot medical VLM claims. We present the work as '
    'hypothesis-generating and emphasise the need for multi-centre external validation; all '
    'conclusions are restricted to the benchmark evaluated.',
    sa=10
)

# --- Body paragraph 5: data, code, ethics, conflicts ---
add_p(
    'The TN5000 dataset is publicly available under its original licence. All code supporting the '
    'benchmark\u2014fine-tuning scripts, inference pipelines, prompt templates, evaluation, and '
    'figure generation\u2014will be released publicly at the repository listed in the Code '
    'Availability section. No new patient data were collected. The authors declare no financial '
    'or non-financial competing interests. No related work is under consideration elsewhere.',
    sa=10
)

# --- Body paragraph 6: authorship ---
add_p(
    'All authors have read and approved the final manuscript and agree to its submission. The '
    'corresponding author\u2019s ORCID and contact details are provided above.',
    sa=12
)

# --- Closing ---
add_p('We thank you for considering our submission.', sa=18)
add_p('Yours sincerely,', sa=18)
add_p('Fahad Wali Ahmed, MD', bold=True, sa=2)
add_p('On behalf of all authors', italic=True, sa=2)

# Save
out_path = '/Users/fahadwali/Downloads/TN5000/SUBMISSION/Main_Manuscript/v4_submission-ready/Cover_Letter.docx'
doc.save(out_path)
print(f'\u2705 Cover letter saved: {out_path}')
